# # extractor.py
# from docx import Document
# import re, os
# from io import BytesIO

# def iter_block_items(parent):
#     from docx.oxml.table import CT_Tbl
#     from docx.oxml.text.paragraph import CT_P
#     from docx.table import Table
#     from docx.text.paragraph import Paragraph

#     for child in parent.element.body:
#         if isinstance(child, CT_P):
#             yield Paragraph(child, parent)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, parent)

# def get_formatted_text(cell):
#     text = ""
#     for paragraph in cell.paragraphs:
#         for run in paragraph.runs:
#             if run.font.subscript:
#                 text += f"<sub>{run.text}</sub>"
#             elif run.font.superscript:
#                 text += f"<sup>{run.text}</sup>"
#             else:
#                 text += run.text
#         text += " "
#     return text.strip()

# def extract_questions_from_docx(file_bytes):
#     doc = Document(BytesIO(file_bytes))
#     units = {}
#     current_unit = "UNIT-I"          # ✅ default fallback unit
#     units[current_unit] = []  

#     for block in iter_block_items(doc):
#         # --- Unit detection ---
#         if block.__class__.__name__ == "Paragraph":
#             text = block.text.strip()
#             match = re.search(r"(UNIT|MODULE)[-\s]*(I|II|III|IV|V)", text, re.I)
#             if match:
#                 current_unit = f"UNIT-{match.group(2).upper()}"
#                 units.setdefault(current_unit, [])
#                 continue

#         # --- Table processing ---
#         if block.__class__.__name__ == "Table":
#             prev_question = None

#             for row in block.rows:
#                 cells = [c.text.strip() for c in row.cells]
#                 if not any(cells):
#                     continue

#                 first_cell = cells[0]
#                 is_new = bool(re.match(r"^\d+[\).]?$", first_cell))

#                 question_text = get_formatted_text(row.cells[1]) if len(row.cells) > 1 else ""

#                 if not is_new and prev_question:
#                     prev_question["text"] += " " + question_text
#                     continue

#                 if current_unit and is_new:
#                     q = {
#                         "text": question_text,
#                         "images": [],
#                         "equations": [],
#                         "tables": []
#                     }
#                     units[current_unit].append(q)
#                     prev_question = q

#     return units


# extractor.py
from docx import Document
import re, os
from io import BytesIO

from statistics import mean
from docx import Document
from io import BytesIO
import re

NUMERIC_RE = re.compile(r"^\d+[\).]?$")

def validate_question_bank_docx(file_bytes):
    """
    Backend-safe validator.
    Accepts DOCX as bytes (from UploadFile.read()).
    """

    doc = Document(BytesIO(file_bytes))
    found_valid_table = False

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        col_count = len(rows[0].cells)
        columns = [[] for _ in range(col_count)]

        for row in rows:
            for i, cell in enumerate(row.cells):
                columns[i].append(cell.text.strip())

        # ---- 1️⃣ Ghost column detection ----
        for idx, col in enumerate(columns):
            empty_ratio = sum(1 for c in col if not c) / len(col)
            if empty_ratio >= 0.8:
                return False, (
                    f"Invalid structure: Ghost / hidden column detected at column {idx+1}. "
                    "Please remove merged or empty columns."
                )

        # ---- 2️⃣ Detect numeric column ----
        numeric_col = None
        for i, col in enumerate(columns):
            ratio = sum(1 for c in col if NUMERIC_RE.match(c)) / len(col)
            if ratio >= 0.6:
                numeric_col = i
                break

        if numeric_col is None:
            return False, "No numeric question number column detected."

        # ---- 3️⃣ Detect question column (longest text) ----
        avg_lengths = []
        for col in columns:
            lengths = [len(c) for c in col if c and not NUMERIC_RE.match(c)]
            avg_lengths.append(mean(lengths) if lengths else 0)

        question_col = avg_lengths.index(max(avg_lengths))

        if question_col == numeric_col:
            return False, "Question column could not be identified."

        meaningful = sum(
            1 for c in columns[question_col] if len(c) >= 15
        ) / len(columns[question_col])

        if meaningful < 0.5:
            return False, "Question column does not contain enough meaningful text."

        found_valid_table = True

    if not found_valid_table:
        return False, "No valid question table found."

    return True, None


def iter_block_items(parent):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def get_formatted_text(cell):
    text = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.subscript:
                text += f"<sub>{run.text}</sub>"
            elif run.font.superscript:
                text += f"<sup>{run.text}</sup>"
            else:
                text += run.text
        text += " "
    return text.strip()

def extract_images_from_cell(cell, unit, q_index, save_dir):
    """Extract and save images from a table cell."""
    images = []
    for paragraph in cell.paragraphs:
        for blip in paragraph._element.xpath(".//a:blip"):
            embed = blip.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            image_part = paragraph.part.related_parts[embed]
            image_data = image_part.blob
            image_name = f"{unit}_Q{q_index}_{len(images)+1}.png"
            image_path = os.path.join(save_dir, image_name)
            with open(image_path, "wb") as f:
                f.write(image_data)
            images.append(image_name)
    return images

def extract_equations_from_cell(cell):
    """Extract math equations (simple text form)."""
    equations = []
    for paragraph in cell.paragraphs:
        for math in paragraph._element.xpath(".//m:oMath | .//m:oMathPara"):
            equations.append(math.text or "[Equation]")
    return equations

def extract_tables_from_cell(cell, unit="TABLE", q_index=0, save_dir="images"):
    """
    Extract nested tables and include text, images, and equations within each cell.
    """
    tables_html = []
    NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    for tbl_index, tbl in enumerate(cell._element.findall(f".//{NS}tbl"), start=1):
        html = "<table border='1'>"

        for tr in tbl.findall(f".//{NS}tr"):
            html += "<tr>"
            for tc_index, tc in enumerate(tr.findall(f".//{NS}tc"), start=1):
                html += "<td>"

                # Extract content from paragraphs inside each table cell
                paragraphs = tc.findall(f".//{NS}p")
                for p_index, p in enumerate(paragraphs, start=1):
                    # --- Extract text ---
                    texts = [t.text for t in p.findall(f".//{NS}t") if t.text]
                    if texts:
                        html += " ".join(texts) + " "

                    # --- Extract images ---
                    for blip in p.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
                        embed = blip.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if not embed:
                            continue
                        image_part = cell.part.related_parts[embed]
                        image_data = image_part.blob
                        img_name = f"{unit}_Q{q_index}_T{tbl_index}_C{tc_index}_{p_index}.png"
                        img_path = os.path.join(save_dir, img_name)
                        with open(img_path, "wb") as f:
                            f.write(image_data)
                        html += f"<img src='{img_name}' alt='image'> "

                    # --- Extract equations ---
                    for math in p.findall(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"):
                        html += "[Equation] "

                html += "</td>"
            html += "</tr>"

        html += "</table>"
        tables_html.append(html)

    return tables_html


def extract_questions_from_docx(file_bytes, image_output_dir="images"):
    if not os.path.exists(image_output_dir):
        os.makedirs(image_output_dir)

    # 🔥 ONLY CHANGE: use BytesIO
    doc = Document(BytesIO(file_bytes))

    units = {}
    current_unit = None

    prev_question = None
    prev_qno = None

    def normalize(text):
        text = text.replace("\t", " ")
        text = re.sub(r"\s+", " ", text)
        return text.strip()
    ROMAN_MAP = {
        "1": "I",
        "2": "II",
        "3": "III",
        "4": "IV",
        "5": "V",
        "I": "I",
        "II": "II",
        "III": "III",
        "IV": "IV",
        "V": "V",
    }

    def detect_unit(text: str):
        if not text:
            return None

        # normalize weird dashes and spaces
        t = text.upper()
        t = t.replace("–", "-").replace("—", "-")
        t = re.sub(r"\s+", " ", t).strip()

        # match UNIT / MODULE with number or roman
        m = re.search(
            r"\b(UNIT|MODULE)\s*[-:]?\s*(\d+|I|II|III|IV|V)\b",
            t
        )

        if not m:
            return None

        raw = m.group(2)
        normalized = ROMAN_MAP.get(raw)

        if not normalized:
            return None

        return f"UNIT-{normalized}"


    def is_qno(text):
        return bool(re.match(r"^\d+[\).]?$", text))

    for block in iter_block_items(doc):

        # ---------- PARAGRAPHS (UNIT HEADINGS) ----------
        if block.__class__.__name__ == "Paragraph":
            text = normalize(block.text)
            if not text:
                continue

            unit = detect_unit(text)
            if unit:
                current_unit = unit
                units.setdefault(current_unit, [])
                prev_question = None
                prev_qno = None
            continue

        # ---------- TABLES ----------
        if block.__class__.__name__ == "Table":
            for row in block.rows:
                cells = [normalize(c.text) for c in row.cells]

                if not any(cells):
                    continue

                combined = " ".join(cells)

                # UNIT inside table
                unit = detect_unit(combined)
                if unit:
                    current_unit = unit
                    units.setdefault(current_unit, [])
                    prev_question = None
                    prev_qno = None
                    continue

                # Skip header rows
                header_words = {"s.no", "question", "co", "bl", "marks"}
                if all(c.lower() in header_words for c in cells if c):
                    continue

                # ---------- FIND QUESTION CELL ----------
                non_empty = [(i, c) for i, c in enumerate(cells) if c]
                if not non_empty:
                    continue

                q_idx, question_text = max(non_empty, key=lambda x: len(x[1]))
                question_text = normalize(question_text)

                # ---------- FIND QUESTION NUMBER ----------
                qno = None
                if q_idx > 0:
                    candidate = cells[q_idx - 1]
                    if is_qno(candidate):
                        qno = candidate

                # ---------- CONTINUATION ----------
                if qno is None:
                    if prev_question:
                        prev_question["text"] += " " + question_text
                    continue

                # ---------- SAME QUESTION NUMBER ----------
                if prev_question and prev_qno == qno:
                    prev_question["text"] += " " + question_text
                    continue

                # ---------- NEW QUESTION ----------
                if current_unit:
                    q = {
                        "qno": qno,
                        "text": question_text,
                        "images": [],
                        "equations": [],
                        "tables": []
                    }

                    for cell in row.cells:
                        q["images"].extend(
                            extract_images_from_cell(
                                cell,
                                current_unit,
                                len(units[current_unit]) + 1,
                                image_output_dir
                            )
                        )
                        q["equations"].extend(extract_equations_from_cell(cell))
                        q["tables"].extend(extract_tables_from_cell(cell))

                    units[current_unit].append(q)
                    prev_question = q
                    prev_qno = qno


    return units
